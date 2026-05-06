"""Multi-format document extractors.

Supports PDF (text + OCR fallback), images (OCR), TXT, Markdown, DOCX, CSV/XLSX.
OCR is loaded lazily via ModelManager.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


class ExtractedDocument:
    """Result of extracting text from a file."""

    def __init__(
        self,
        text: str,
        pages: int = 1,
        file_type: str = "",
        ocr_used: bool = False,
        metadata: dict[str, Any] | None = None,
    ) -> None:
        self.text = text
        self.pages = pages
        self.file_type = file_type
        self.ocr_used = ocr_used
        self.metadata = metadata or {}


def extract_file(
    path: Path,
    ocr_reader: Any | None = None,
    ocr_enabled: bool = True,
    ocr_languages: list[str] | None = None,
) -> ExtractedDocument:
    """Route to the correct extractor based on file extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path, ocr_reader, ocr_enabled, ocr_languages)
    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"):
        return extract_image(path, ocr_reader, ocr_enabled, ocr_languages)
    if ext in (".txt", ".md", ".markdown"):
        return extract_text(path, ext)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".csv", ".xlsx", ".xls"):
        return extract_spreadsheet(path, ext)

    logger.warning("unsupported_extension", extra={"path": str(path), "ext": ext})
    return ExtractedDocument(text="", file_type=ext, metadata={"error": "unsupported_extension"})


# ------------------------------------------------------------------
# PDF
# ------------------------------------------------------------------

def extract_pdf(
    path: Path,
    ocr_reader: Any | None = None,
    ocr_enabled: bool = False,
    ocr_languages: list[str] | None = None,
) -> ExtractedDocument:
    import fitz  # PyMuPDF

    ocr_used = False
    pages_text = []

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        logger.error("pdf_open_failed", extra={"path": str(path), "error": str(exc)})
        return ExtractedDocument(text="", file_type=".pdf", metadata={"error": "pdf_open_failed"})

    low_text_pages = 0
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        # Heuristic: page with very little text is likely scanned
        if len(text.strip()) >= 20:
            pages_text.append(f"[Page {page_num + 1}]\n{text}")
        else:
            low_text_pages += 1
            if ocr_enabled and model_manager:
                try:
                    pix = page.get_pixmap(dpi=200)
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        tmp_path = Path(tmp.name)
                        pix.save(str(tmp_path))
                    ocr_result = extract_image(tmp_path, ocr_reader, ocr_enabled, ocr_languages)
                    if ocr_result.text.strip():
                        pages_text.append(f"[Page {page_num + 1} (OCR)]\n{ocr_result.text}")
                        ocr_used = True
                    tmp_path.unlink(missing_ok=True)
                except Exception as exc:
                    logger.warning("pdf_ocr_page_failed", extra={"page": page_num, "error": str(exc)})

    doc.close()
    full_text = "\n\n".join(pages_text)
    return ExtractedDocument(
        text=full_text,
        pages=len(doc),
        file_type=".pdf",
        ocr_used=ocr_used,
        metadata={"low_text_pages": low_text_pages},
    )


# ------------------------------------------------------------------
# Images (OCR)
# ------------------------------------------------------------------

def extract_image(
    path: Path,
    ocr_reader: Any | None,
    ocr_enabled: bool,
    ocr_languages: list[str] | None,
) -> ExtractedDocument:
    if not ocr_enabled:
        return ExtractedDocument(text="", file_type=path.suffix.lower(), metadata={"error": "ocr_disabled"})

    reader = ocr_reader
    if reader is None:
        try:
            import easyocr
            reader = easyocr.Reader(ocr_languages or ["fra", "eng"], gpu=False)
        except Exception as exc:
            logger.error("ocr_init_failed", extra={"path": str(path), "error": str(exc)})
            return ExtractedDocument(text="", file_type=path.suffix.lower(), metadata={"error": "ocr_init_failed"})

    try:
        results = reader.readtext(str(path), detail=0)
        text = "\n".join(results)
    except Exception as exc:
        logger.error("ocr_failed", extra={"path": str(path), "error": str(exc)})
        text = ""

    return ExtractedDocument(
        text=text,
        pages=1,
        file_type=path.suffix.lower(),
        ocr_used=True,
    )


# ------------------------------------------------------------------
# Plain text / Markdown
# ------------------------------------------------------------------

def extract_text(path: Path, ext: str) -> ExtractedDocument:
    try:
        raw = path.read_bytes()
        try:
            content = raw.decode("utf-8")
        except UnicodeDecodeError:
            content = raw.decode("latin-1", errors="replace")
        return ExtractedDocument(text=content, pages=1, file_type=ext)
    except Exception as exc:
        logger.error("text_read_failed", extra={"path": str(path), "error": str(exc)})
        return ExtractedDocument(text="", file_type=ext, metadata={"error": str(exc)})


# ------------------------------------------------------------------
# DOCX
# ------------------------------------------------------------------

def extract_docx(path: Path) -> ExtractedDocument:
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n\n".join(paragraphs)
        return ExtractedDocument(text=text, pages=len(doc.paragraphs) // 20 + 1, file_type=".docx")
    except Exception as exc:
        logger.error("docx_extract_failed", extra={"path": str(path), "error": str(exc)})
        return ExtractedDocument(text="", file_type=".docx", metadata={"error": str(exc)})


# ------------------------------------------------------------------
# Spreadsheet
# ------------------------------------------------------------------

def extract_spreadsheet(path: Path, ext: str) -> ExtractedDocument:
    try:
        import pandas as pd

        if ext == ".csv":
            df = pd.read_csv(str(path), encoding="utf-8")
        else:
            df = pd.read_excel(str(path))

        # Serialize rows into lines of text
        lines = []
        for _, row in df.iterrows():
            row_text = " | ".join(str(v) for v in row.values if pd.notna(v))
            if row_text:
                lines.append(row_text)
        text = "\n\n".join(lines)
        return ExtractedDocument(text=text, pages=1, file_type=ext)
    except Exception as exc:
        logger.error("spreadsheet_extract_failed", extra={"path": str(path), "error": str(exc)})
        return ExtractedDocument(text="", file_type=ext, metadata={"error": str(exc)})


# ------------------------------------------------------------------
# Cleaning helpers
# ------------------------------------------------------------------

def clean_text(text: str) -> str:
    """Normalize and strip noise."""
    # Replace non-breaking spaces, strip Unicode artifacts
    import unicodedata

    text = unicodedata.normalize("NFKC", text)
    # Collapse multiple spaces and blank lines
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()
